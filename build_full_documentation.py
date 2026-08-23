import os
import sys
import json
import shutil
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOCS_DIR = Path("docs")
DOCS_DIR.mkdir(parents=True, exist_ok=True)
DOCX_OUT = DOCS_DIR / "Telugu_HCR_v4_Project_Documentation.docx"
MD_OUT = DOCS_DIR / "PROJECT_DOCUMENTATION.md"
DOWNLOADS_DIR = Path.home() / "Downloads"

doc = docx.Document()

# Configure normal style
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Cambria'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.font.name = 'Cambria'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)
    return p

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Cambria'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Cambria'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x57, 0x42, 0x39)
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Cambria'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xA0, 0x3D, 0x00)
    return p

def add_paragraph(text, bold_prefix=None, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = 'Cambria'
        r_b.font.size = Pt(11)
        r_b.font.bold = True
        r_b.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)
    if text:
        r_t = p.add_run(text)
        r_t.font.name = 'Cambria'
        r_t.font.size = Pt(11)
        r_t.font.italic = italic
        r_t.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = 'Cambria'
        r_b.font.size = Pt(11)
        r_b.font.bold = True
        r_b.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)
    if text:
        r_t = p.add_run(text)
        r_t.font.name = 'Cambria'
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)
    return p

def add_code_block(code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F5F3EE")
    set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1B)
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)

def add_terminal_command(cmd_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "EDEBE6")
    set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(cmd_text.strip())
    run.font.name = 'Consolas'
    run.font.size = Pt(9.0)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xA0, 0x3D, 0x00)
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(3)

def add_image_figure(image_path, caption_text, width_inches=5.8):
    if not Path(image_path).exists():
        print(f"Warning: Image {image_path} not found.")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(2)
    p_img.paragraph_format.keep_with_next = True
    r = p_img.add_run()
    r.add_picture(str(image_path), width=Inches(width_inches))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(8)
    r_cap = p_cap.add_run(f"Figure: {caption_text}")
    r_cap.font.name = 'Cambria'
    r_cap.font.size = Pt(9.5)
    r_cap.font.italic = True
    r_cap.font.color.rgb = RGBColor(0x57, 0x42, 0x39)

print("Starting document creation...")

# -------------------------------------------------------------
# Document Header and Description
# -------------------------------------------------------------
add_title("Telugu Handwritten Character Recognizer (Telugu HCR v4)")

add_paragraph(
    "Telugu is a Dravidian language with an abugida writing system where characters (aksharas) are formed through structural combinations of base consonants, secondary vowel signs (guninthamulu), and subscript conjunct consonants (vattulu). The Telugu dataset contains 630 character directories, which collapse to 596 structurally unique compound graphemes across 34 confirmed duplicate pairs between the base consonant and guninthamulu folders. Monolithic 630-way flat classification models struggle on this task because data is heavily fragmented across hundreds of compound classes, and visual features share sub-character components."
)

add_paragraph(
    "Telugu HCR v4 addresses this by decomposing every compound character into three linguistic components: a Base Akshara (52 classes), a Vowel Modifier (16 classes), and a Subscript Conjunct Vattu (37 classes). An EfficientNetV2-B0 backbone extracts unified visual representations, while three independent dense heads predict the component distributions simultaneously. A Constrained Maximum Likelihood Decoding engine recombines these distributions over the 596 valid compound combinations, achieving 94.72% Top-1 and 99.29% Top-5 recombined accuracy on the test set. An interactive Streamlit web application with stroke-level Gaussian blur preprocessing provides real-time handwriting recognition."
)

# -------------------------------------------------------------
# Scenarios
# -------------------------------------------------------------
add_heading1("Scenarios")

add_paragraph(
    "A user draws a complex conjunct character such as 'ksha' (క్ష) or 'thri' (త్రి) on the interactive drawing canvas. The application crops and centers the glyph, applies Gaussian blur smoothing to bridge the domain shift between synthetic digital canvas strokes and paper ink dispersion, runs multi-head inference, and presents the predicted Telugu Unicode character along with decomposed component confidence scores.",
    bold_prefix="Scenario 1: Interactive Digitization of Complex Compound Graphemes: "
)

add_paragraph(
    "An archival institution scans physical historical Telugu documents and palm-leaf manuscripts. The batch evaluation pipeline reads cropped character regions, normalizes them through the single source of truth preprocessing module, and outputs structured predictions with Top-5 candidate rankings, enabling automated transcription and human-in-the-loop verification.",
    bold_prefix="Scenario 2: Historical Manuscript Archival and Document OCR: "
)

add_paragraph(
    "A student learning the Telugu script writes aksharas on a tablet device. The system evaluates stroke structure and provides immediate visual feedback, showing whether mistakes occurred in the base letter, the vowel matra, or the subscript vattu, assisting learners in mastering complex handwriting orthography.",
    bold_prefix="Scenario 3: Language Learning and Calligraphic Feedback: "
)

# -------------------------------------------------------------
# Technical Architecture
# -------------------------------------------------------------
add_heading1("Technical Architecture:")

add_image_figure(
    "outputs/docs_diagrams/architecture_diagram.png",
    "Telugu HCR v4 Multi-Head EfficientNetV2 Architecture with Constrained MLE Decoding"
)

add_paragraph(
    "The technical architecture consists of four interconnected subsystems: the data preprocessing pipeline, the multi-head neural network backbone, the constrained decoding engine, and the Streamlit user interface."
)

add_bullet("Single Source of Truth Preprocessing: Decodes raw image bytes or arrays, extracts single-channel grayscale, applies aspect-ratio preserving padding to square with white background (255.0), resizes to 128x128 pixels using bilinear interpolation, replicates to 3 channels, and normalizes according to EfficientNetV2 scaling.", bold_prefix="Input Pipeline: ")

add_bullet("EfficientNetV2-B0 backbone pretrained on ImageNet (5.9M parameters), followed by Global Average Pooling, Batch Normalization, and Dropout (0.3).", bold_prefix="Backbone Network: ")

add_bullet("Three parallel dense branches: Base Head (Dense 256, Dropout 0.3, Softmax 52), Modifier Head (Dense 128, Dropout 0.3, Softmax 16), and Vattu Head (Dense 128, Dropout 0.3, Softmax 37).", bold_prefix="Multi-Task Classification Heads: ")

add_bullet("Computes greedy argmax combination, evaluates validity against label maps, and applies Constrained Maximum Likelihood Decoding over 596 valid triples with log-likelihood optimization if the greedy combination is invalid.", bold_prefix="Recombination Engine: ")

add_bullet("Streamlit web application with a 2-column matte utility layout, stroke smoothing via Gaussian blur (sigma=3.0), and real-time Telugu Unicode glyph assembly.", bold_prefix="Interactive Frontend: ")

# -------------------------------------------------------------
# Pre-requisites
# -------------------------------------------------------------
add_heading1("Pre-requisites:")

add_bullet("Python 3.10 or higher (tested with Python 3.10.13 via pyenv).", bold_prefix="Python Environment: ")
add_bullet("TensorFlow 2.16+ and Keras 3.0+ for multi-head model architecture, AdamW optimizer, EMA weight tracking, and tf.data pipelines.", bold_prefix="Deep Learning Frameworks: ")
add_bullet("Streamlit 1.30+ and streamlit-drawable-canvas 0.9.0+ for the browser-based drawing interface and prediction display.", bold_prefix="Web Application: ")
add_bullet("OpenCV (opencv-python), NumPy, Pandas, Scikit-Learn, Pillow, PyYAML, and Tqdm for image processing, metrics calculation, and dataset handling.", bold_prefix="Supporting Libraries: ")
add_bullet("Directory Test1 containing Achulu, Hallulu, Guninthamulu, and Othulu folders across 630 class directories.", bold_prefix="Telugu Dataset: ")
add_bullet("Apple Silicon Mac (MPS / CPU) or Linux / Windows workstation with NVIDIA CUDA GPU.", bold_prefix="Hardware: ")

# -------------------------------------------------------------
# Project Workflow Overview
# -------------------------------------------------------------
add_heading1("Project Workflow:")

# Table of Milestones
table_wf = doc.add_table(rows=7, cols=2)
table_wf.alignment = WD_TABLE_ALIGNMENT.CENTER
table_wf.autofit = False

headers = ["Milestone", "Activities & Core Focus"]
for i, h in enumerate(headers):
    cell = table_wf.cell(0, i)
    set_cell_background(cell, "EDEBE6")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True
    r.font.size = Pt(10)

wf_data = [
    ("Milestone 1: Data Preparation & Architecture Definition", "Activity 1.1: Dataset structure & stratified splitting (80/10/10)\nActivity 1.2: Class collision & duplicate pairs audit (34 pairs)\nActivity 1.3: Single source of truth preprocessing pipeline\nActivity 1.4: Multi-head grapheme decomposition engine"),
    ("Milestone 2: Model Architecture & Loss Design", "Activity 2.1: Multi-head EfficientNetV2 backbone construction\nActivity 2.2: Class-weighted loss with label smoothing\nActivity 2.3: Data augmentation & multi-head CutMix pipeline"),
    ("Milestone 3: Training Pipeline & Checkpointing", "Activity 3.1: Two-phase training protocol (warmup + fine-tuning)\nActivity 3.2: Early timing extrapolation & Kaggle session guard\nActivity 3.3: Atomic full-state checkpoint management"),
    ("Milestone 4: Model Evaluation & Analysis", "Activity 4.1: Multi-head and recombined test set evaluation\nActivity 4.2: Baseline benchmarking & confusion pair analysis"),
    ("Milestone 5: Streamlit Web Application", "Activity 5.1: Two-column matte technical utility interface\nActivity 5.2: Canvas preprocessing & stroke smoothing\nActivity 5.3: Top-3 prediction cards & Unicode glyph assembly"),
    ("Milestone 6: Deployment & Verification", "Activity 6.1: Local virtual environment & dependency setup\nActivity 6.2: Local testing & canvas verification\nActivity 6.3: Public deployment via Ngrok & sharing")
]

for row_idx, (m_title, m_desc) in enumerate(wf_data, start=1):
    c0 = table_wf.cell(row_idx, 0)
    c1 = table_wf.cell(row_idx, 1)
    c0.width = Inches(2.3)
    c1.width = Inches(4.2)
    set_cell_margins(c0, top=80, bottom=80, left=120, right=120)
    set_cell_margins(c1, top=80, bottom=80, left=120, right=120)
    if row_idx % 2 == 1:
        set_cell_background(c0, "FBF9F4")
        set_cell_background(c1, "FBF9F4")
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(m_title)
    r0.font.bold = True
    r0.font.size = Pt(9.5)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(m_desc)
    r1.font.size = Pt(9.0)

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(6)

# -------------------------------------------------------------
# Milestone 1: Data Preparation and Architecture Definition
# -------------------------------------------------------------
add_heading1("Milestone 1: Data Preparation and Architecture Definition")

add_paragraph(
    "Milestone 1 establishes the foundational data architecture for Telugu HCR v4. This involves auditing the dataset structure, identifying semantic class duplicates, implementing a centralized single source of truth preprocessing pipeline, and creating the multi-head grapheme decomposition engine that maps compound Telugu characters into linguistic primitives."
)

add_heading2("Activity 1.1: Dataset Structure and Stratified Splitting")

add_paragraph(
    "Step 1: Inspect Dataset Hierarchy: The raw dataset consists of four main categories: Achulu (vowels, 16 folders), Hallulu (consonants, 36 folders), Guninthamulu (vowel-consonant combinations, 542 folders across sub-directories), and Othulu (subscript vattu markers, 36 folders), yielding 630 total class directories."
)

add_paragraph(
    "Step 2: Generate Stratified Splits: An 80/10/10 stratified split is created using a fixed random seed (42) to ensure reproducibility. Each class directory is sampled proportionally across train.csv, val.csv, and test.csv."
)

add_paragraph(
    "Step 3: Export Split Manifests and Label Maps: The script parses all image paths, generates numeric label indices for base akshara, vowel modifier, and conjunct vattu, and writes outputs/label_maps.json."
)

add_paragraph("Command to generate stratified splits and label maps:")
add_terminal_command("python src/data/split.py --dataset_root \"data/Test1\" --output_dir \"outputs\"")

add_paragraph("Source Code: src/data/split.py (Dataset Collection and Split Generation):")
add_code_block("""def collect_dataset_images(dataset_root: str, relative_paths: bool = True) -> Tuple[List[Dict[str, Any]], List[str]]:
    root = Path(dataset_root)
    if not root.exists():
        raise FileNotFoundError(f"Dataset root '{dataset_root}' does not exist.")
    records = []
    class_set = set()
    valid_extensions = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"}
    for cat_dir in sorted(root.iterdir()):
        if not cat_dir.is_dir() or cat_dir.name.startswith("."):
            continue
        cat = cat_dir.name
        if cat.lower() == "guninthamulu":
            for c_dir in sorted(cat_dir.iterdir()):
                if not c_dir.is_dir() or c_dir.name.startswith("."):
                    continue
                for v_dir in sorted(c_dir.iterdir()):
                    if not v_dir.is_dir() or v_dir.name.startswith("."):
                        continue
                    class_name = f"{cat}__{c_dir.name}__{v_dir.name}"
                    class_set.add(class_name)
                    for img_file in sorted(v_dir.iterdir()):
                        if img_file.is_file() and img_file.suffix.lower() in valid_extensions:
                            fp = str(img_file.relative_to(root)) if relative_paths else str(img_file.resolve())
                            records.append({"file_path": fp, "class_name": class_name, "category": cat})
        else:
            for c_dir in sorted(cat_dir.iterdir()):
                if not c_dir.is_dir() or c_dir.name.startswith("."):
                    continue
                class_name = f"{cat}__{c_dir.name}"
                class_set.add(class_name)
                for img_file in sorted(c_dir.iterdir()):
                    if img_file.is_file() and img_file.suffix.lower() in valid_extensions:
                        fp = str(img_file.relative_to(root)) if relative_paths else str(img_file.resolve())
                        records.append({"file_path": fp, "class_name": class_name, "category": cat})
    return records, sorted(list(class_set))""")

add_heading2("Activity 1.2: Class Collision and Semantic Equivalence Audit")

add_paragraph(
    "Step 1: Audit Confirmed Duplicate Class Pairs: Careful visual and linguistic inspection revealed that 34 class directories in Guninthamulu (specifically the 'none' modifier sub-folder for each consonant) contain the exact same base consonant glyph as the corresponding bare consonant folder in Hallulu (Pattern A duplicates). For example, 'Guninthamulu__kha__ka' and 'hallulu__ka' represent the identical character 'క' (ka)."
)

add_paragraph(
    "Step 2: Canonical Equivalence Grouping: Rather than treating these 34 pairs as distinct competing classes (which causes artificial confusion), src/data/known_duplicates.py maintains an equivalence group lookup that collapses 630 class folders to 596 unique compound combinations during evaluation and recombination."
)

add_image_figure(
    "outputs/collision_review/pattern_a_hallulu_b_vs_guninthamulu_ba_b.png",
    "Sample Class Collision Audit: Confirmed Identical Character Pair (hallulu__b vs Guninthamulu__ba__b)"
)

add_paragraph("Source Code: src/data/known_duplicates.py (Duplicate Resolution Engine):")
add_code_block("""CONFIRMED_DUPLICATE_CLASSES: List[Tuple[str, str]] = [
    ("Guninthamulu__kha__ka", "hallulu__ka"),
    ("Guninthamulu__khh__kha", "hallulu__kha"),
    ("Guninthamulu__ga__ga", "hallulu__g"),
    ("Guninthamulu__gha__gha", "hallulu__gh"),
    ("Guninthamulu__ch__ch", "hallulu__ch"),
    ("Guninthamulu__cha__ch", "hallulu__cha"),
    ("Guninthamulu__ja__j", "hallulu__jh"),
    ("Guninthamulu__jh__jh", "hallulu__jha"),
    ("Guninthamulu__ta__ta", "hallulu__ta"),
    ("Guninthamulu__tt__t", "hallulu__th"),
    ("Guninthamulu__d__d", "hallulu__d"),
    ("Guninthamulu__dh__dh", "hallulu__dh"),
    ("Guninthamulu__ana__an", "hallulu__ana"),
    ("Guninthamulu__th__th", "hallulu__tha"),
    ("Guninthamulu__tha__th", "hallulu__thah"),
    ("Guninthamulu__da__da", "hallulu__da"),
    ("Guninthamulu__dha__dh", "hallulu__dha"),
    ("Guninthamulu__na__n", "hallulu__n"),
    ("Guninthamulu__pa__p", "hallulu__P"),
    ("Guninthamulu__pha__p", "hallulu__Ph"),
    ("Guninthamulu__ba__b", "hallulu__b"),
    ("Guninthamulu__bha__bh", "hallulu__bh"),
    ("Guninthamulu__ma__m", "hallulu__m"),
    ("Guninthamulu__ya__y", "hallulu__y"),
    ("Guninthamulu__ra__r", "hallulu__r"),
    ("Guninthamulu__RR__rr", "hallulu__rr"),
    ("Guninthamulu__l__l", "hallulu__l"),
    ("Guninthamulu__ll__l", "hallulu__ll"),
    ("Guninthamulu__va__v", "hallulu__v"),
    ("Guninthamulu__sha__sh", "hallulu__s"),
    ("Guninthamulu__sh__sh", "hallulu__sh"),
    ("Guninthamulu__sa__s", "hallulu__sa"),
    ("Guninthamulu__ha__h", "hallulu__h"),
    ("Guninthamulu__ksh__ksh", "hallulu__ks"),
]

_CANONICAL_LOOKUP: Dict[str, str] = {}
_EQUIVALENCE_GROUPS: Dict[str, Set[str]] = {}
for c1, c2 in CONFIRMED_DUPLICATE_CLASSES:
    canonical = min(c1, c2)
    _CANONICAL_LOOKUP[c1] = canonical
    _CANONICAL_LOOKUP[c2] = canonical
    group = {c1, c2}
    _EQUIVALENCE_GROUPS[c1] = group
    _EQUIVALENCE_GROUPS[c2] = group

def get_canonical_class_name(class_name: str) -> str:
    return _CANONICAL_LOOKUP.get(class_name, class_name)

def is_known_duplicate_pair(class1: str, class2: str) -> bool:
    return class1 in _EQUIVALENCE_GROUPS and class2 in _EQUIVALENCE_GROUPS[class1]""")

add_heading2("Activity 1.3: Single Source of Truth Image Preprocessing")

add_paragraph(
    "Step 1: Pad to Square with White Background: Handwritten character crops have varying aspect ratios. Direct non-uniform resizing distorts character strokes and aspect ratios. The pad_to_square function computes the maximum dimension and pads symmetrically using constant value 255.0 (white paper background)."
)

add_paragraph(
    "Step 2: Grayscale Extraction and EfficientNetV2 Normalization: Multichannel inputs (RGB or RGBA) are converted to single-channel grayscale via standard luminance weights, resized to 128x128 using bilinear interpolation, replicated across 3 channels, and scaled using EfficientNetV2 preprocess_input."
)

add_image_figure(
    "outputs/docs_diagrams/preprocessing_pipeline.png",
    "Single Source of Truth Image Preprocessing Pipeline"
)

add_paragraph("Source Code: src/data/preprocessing.py (Complete Module):")
add_code_block("""from typing import Union
import numpy as np
import tensorflow as tf
from tensorflow.keras.applications.efficientnet_v2 import preprocess_input

IMAGE_SIZE: int = 128
NUM_CHANNELS: int = 3
_SAMPLE_WHITE = tf.constant([[[255.0, 255.0, 255.0]]], dtype=tf.float32)
BACKGROUND_FILL_VALUE: float = float(preprocess_input(_SAMPLE_WHITE)[0, 0, 0].numpy())

def pad_to_square(image: tf.Tensor, pad_value: float = 255.0) -> tf.Tensor:
    shape = tf.shape(image)
    h, w = shape[0], shape[1]
    max_dim = tf.maximum(h, w)
    pad_h, pad_w = max_dim - h, max_dim - w
    pad_top, pad_bottom = pad_h // 2, pad_h - (pad_h // 2)
    pad_left, pad_right = pad_w // 2, pad_w - (pad_w // 2)
    image = tf.cond(tf.equal(tf.rank(image), 2), lambda: tf.expand_dims(image, axis=-1), lambda: image)
    paddings = [[pad_top, pad_bottom], [pad_left, pad_right], [0, 0]]
    return tf.pad(image, paddings, mode="CONSTANT", constant_values=pad_value)

def preprocess_image(raw_image_bytes_or_array: Union[tf.Tensor, np.ndarray, bytes], img_size: int = IMAGE_SIZE) -> tf.Tensor:
    if isinstance(raw_image_bytes_or_array, (bytes, bytearray)):
        img = tf.io.decode_image(raw_image_bytes_or_array, channels=0, expand_animations=False)
        img = tf.cast(img, tf.float32)
        img = _ensure_single_channel_grayscale(img)
    elif isinstance(raw_image_bytes_or_array, tf.Tensor) and raw_image_bytes_or_array.dtype == tf.string:
        img = tf.io.decode_image(raw_image_bytes_or_array, channels=0, expand_animations=False)
        img = tf.cast(img, tf.float32)
        img = _ensure_single_channel_grayscale(img)
    else:
        tensor = tf.convert_to_tensor(raw_image_bytes_or_array)
        img = tf.cast(tensor, tf.float32)
        img = _ensure_single_channel_grayscale(img)
    padded = pad_to_square(img, pad_value=255.0)
    resized = tf.image.resize(padded, [img_size, img_size], method=tf.image.ResizeMethod.BILINEAR)
    replicated = tf.repeat(resized, repeats=3, axis=-1)
    return preprocess_input(replicated)""")

add_heading2("Activity 1.4: Multi-Head Decomposition Engine")

add_paragraph(
    "Step 1: Grapheme Primitives Parsing: Every folder class name is parsed by decompose_class_name into (base_letter, vowel_modifier, conjunct_vattu). For example, 'Guninthamulu__ka__kii' decomposes into base consonant 'క' (ka), vowel modifier 'ii' (ఈకారం), and vattu 'none'."
)

add_paragraph(
    "Step 2: Constrained Maximum Likelihood Decoding: During inference, the three classification heads output probability vectors p_base, p_mod, and p_vattu. recombine_prediction performs a fast greedy argmax lookup. If the resulting combination is valid in the dataset, it is returned directly. If the greedy combination is invalid, it evaluates log p_base(b) + log p_mod(m) + log p_vattu(v) over all 596 valid combinations and selects the argmax, guaranteeing 100% valid Telugu grapheme outputs."
)

add_paragraph("Source Code: src/data/decomposition.py (Constrained MLE Decoding):")
add_code_block("""def recombine_prediction(base_probs: np.ndarray, mod_probs: np.ndarray, vattu_probs: np.ndarray, label_maps: Dict[str, Any]) -> Dict[str, Any]:
    base_probs = np.asarray(base_probs, dtype=np.float64)
    mod_probs = np.asarray(mod_probs, dtype=np.float64)
    vattu_probs = np.asarray(vattu_probs, dtype=np.float64)
    eps = 1e-12
    log_p_b = np.log(np.clip(base_probs, eps, 1.0))
    log_p_m = np.log(np.clip(mod_probs, eps, 1.0))
    log_p_v = np.log(np.clip(vattu_probs, eps, 1.0))
    comb_map = label_maps["combination_to_class"]
    valid_triples = label_maps["valid_triples"]
    base_letters = label_maps["base_letters"]
    vowel_modifiers = label_maps["vowel_modifiers"]
    conjunct_modifiers = label_maps["conjunct_modifiers"]
    b_star, m_star, v_star = int(np.argmax(base_probs)), int(np.argmax(mod_probs)), int(np.argmax(vattu_probs))
    greedy_key = f"{b_star}_{m_star}_{v_star}"
    if greedy_key in comb_map:
        pred_class = comb_map[greedy_key]
        pred_b, pred_m, pred_v = b_star, m_star, v_star
        confidence = float(base_probs[b_star] * mod_probs[m_star] * vattu_probs[v_star])
        is_fallback = False
    else:
        best_score = -np.inf
        best_triple = valid_triples[0]
        for b, m, v in valid_triples:
            score = log_p_b[b] + log_p_m[m] + log_p_v[v]
            if score > best_score:
                best_score = score
                best_triple = [b, m, v]
        pred_b, pred_m, pred_v = best_triple
        pred_class = comb_map[f"{pred_b}_{pred_m}_{pred_v}"]
        confidence = float(np.exp(best_score))
        is_fallback = True
    all_scores = []
    total_valid_mass = 0.0
    for b, m, v in valid_triples:
        joint_prob = float(base_probs[b] * mod_probs[m] * vattu_probs[v])
        total_valid_mass += joint_prob
        all_scores.append({"class_name": comb_map[f"{b}_{m}_{v}"], "base_letter": base_letters[b], "vowel_modifier": vowel_modifiers[m], "vattu": conjunct_modifiers[v], "probability": joint_prob})
    if total_valid_mass > 0.0:
        for s in all_scores:
            s["probability"] = float(s["probability"] / total_valid_mass)
    all_scores.sort(key=lambda x: x["probability"], reverse=True)
    return {"predicted_class": pred_class, "base_letter": base_letters[pred_b], "vowel_modifier": vowel_modifiers[pred_m], "vattu": conjunct_modifiers[pred_v], "confidence": confidence, "is_fallback": is_fallback, "top_5": all_scores[:5]}""")

# -------------------------------------------------------------
# Milestone 2: Model Architecture and Loss Design
# -------------------------------------------------------------
add_heading1("Milestone 2: Model Architecture and Loss Design")

add_paragraph(
    "Milestone 2 details the multi-task neural network architecture, custom class-weighted loss formulation, and regularized data augmentation pipeline designed specifically for compound character recognition."
)

add_heading2("Activity 2.1: Multi-Head EfficientNetV2 Backbone Construction")

add_paragraph(
    "Step 1: EfficientNetV2 Feature Extraction: EfficientNetV2-B0 provides an optimal trade-off between model capacity (5.9M parameters), inference speed, and fine-grained feature representation. The network accepts (128, 128, 3) image tensors and produces a 1280-dimensional feature vector after Global Average Pooling."
)

add_paragraph(
    "Step 2: Independent Classification Branches: A Batch Normalization layer and Dropout (0.3) stabilize backbone representations. Three separate dense heads branch off: Base Head (256 units + ReLU + Dropout -> 52 Softmax), Modifier Head (128 units + ReLU + Dropout -> 16 Softmax), and Vattu Head (128 units + ReLU + Dropout -> 37 Softmax)."
)

add_paragraph("Source Code: src/models/multitask_effnetv2.py (Model Architecture Builder):")
add_code_block("""def build_multitask_effnetv2(variant: str = "B0", num_base: int = 52, num_mod: int = 16, num_vattu: int = 37, input_shape: Tuple[int, int, int] = (128, 128, 3), weights: Optional[str] = "imagenet", backbone_trainable: bool = False, dropout_rate: float = 0.3) -> Model:
    inputs = layers.Input(shape=input_shape, name="image_input")
    backbone = EfficientNetV2B0(include_top=False, weights=weights, input_tensor=inputs, pooling=None)
    backbone.trainable = backbone_trainable
    features = backbone.output
    pooled = layers.GlobalAveragePooling2D(name="backbone_gap")(features)
    pooled = layers.BatchNormalization(name="backbone_bn")(pooled)
    pooled = layers.Dropout(dropout_rate, name="backbone_dropout")(pooled)
    base_h = layers.Dense(256, activation="relu", name="base_dense")(pooled)
    base_h = layers.Dropout(dropout_rate, name="base_dropout")(base_h)
    base_out = layers.Dense(num_base, activation="softmax", dtype="float32", name="base_output")(base_h)
    mod_h = layers.Dense(128, activation="relu", name="modifier_dense")(pooled)
    mod_h = layers.Dropout(dropout_rate, name="modifier_dropout")(mod_h)
    mod_out = layers.Dense(num_mod, activation="softmax", dtype="float32", name="modifier_output")(mod_h)
    vattu_h = layers.Dense(128, activation="relu", name="vattu_dense")(pooled)
    vattu_h = layers.Dropout(dropout_rate, name="vattu_dropout")(vattu_h)
    vattu_out = layers.Dense(num_vattu, activation="softmax", dtype="float32", name="vattu_output")(vattu_h)
    return Model(inputs=inputs, outputs=[base_out, mod_out, vattu_out], name=f"multitask_effnetv2_{variant.lower()}")""")

add_heading2("Activity 2.2: Class-Weighted Loss Formulation with Label Smoothing")

add_paragraph(
    "Step 1: Compute Frequency-Balanced Inverse Class Weights: Certain modifiers and vattus occur infrequently in Telugu text. To prevent dominant classes from overpowering gradient updates, compute_normalized_class_weights calculates inverse frequency weights clipped between 0.1 and 10.0 and normalized to mean 1.0."
)

add_paragraph(
    "Step 2: Implement Weighted Categorical Cross-Entropy: Keras standard sample_weight mechanisms can experience instability with multi-head outputs. WeightedCategoricalCrossentropy bakes per-class weight multipliers and label smoothing directly into the tensor loss graph, supporting continuous soft targets from CutMix."
)

add_paragraph("Source Code: src/models/losses.py (Custom Loss Function):")
add_code_block("""@tf.keras.utils.register_keras_serializable(package="TeluguHCR")
class WeightedCategoricalCrossentropy(tf.keras.losses.Loss):
    def __init__(self, class_weights=None, label_smoothing=0.0, from_logits=False, name="weighted_categorical_crossentropy", **kwargs):
        super().__init__(name=name, **kwargs)
        self.label_smoothing = float(label_smoothing)
        self.from_logits = bool(from_logits)
        self.class_weights = tf.constant(class_weights, dtype=tf.float32) if class_weights is not None else None

    def call(self, y_true: tf.Tensor, y_pred: tf.Tensor) -> tf.Tensor:
        y_true = tf.cast(y_true, tf.float32)
        y_pred = tf.cast(y_pred, tf.float32)
        num_classes = tf.cast(tf.shape(y_true)[-1], tf.float32)
        if self.label_smoothing > 0.0:
            y_true = y_true * (1.0 - self.label_smoothing) + (self.label_smoothing / num_classes)
        if self.from_logits:
            y_pred = tf.nn.softmax(y_pred, axis=-1)
        eps = tf.keras.backend.epsilon()
        y_pred_safe = tf.clip_by_value(y_pred, eps, 1.0 - eps)
        ce = -y_true * tf.math.log(y_pred_safe)
        if self.class_weights is not None:
            ce = ce * self.class_weights
        return tf.reduce_mean(tf.reduce_sum(ce, axis=-1))""")

add_heading2("Activity 2.3: Data Augmentation and Multi-Head CutMix Pipeline")

add_paragraph(
    "Step 1: Spatial Augmentations: Spatial transformations include RandomRotation (+/-5 degrees), RandomTranslation (+/-5% height and width), and RandomZoom (+/-5%). All boundary padding explicitly utilizes BACKGROUND_FILL_VALUE from preprocessing.py."
)

add_paragraph(
    "Step 2: Multi-Head CutMix Blending: CutMix cuts a rectangular patch from a secondary image and pastes it onto the primary image. The area ratio lambda adjusts the one-hot target vectors across all three classification heads simultaneously, improving generalization."
)

add_paragraph("Source Code: src/data/augmentation.py (CutMix Implementation):")
add_code_block("""@tf.function
def apply_cutmix(images: tf.Tensor, base_labels: tf.Tensor, mod_labels: tf.Tensor, vattu_labels: tf.Tensor, alpha: float = 0.4):
    batch_size = tf.shape(images)[0]
    h, w = tf.cast(tf.shape(images)[1], tf.float32), tf.cast(tf.shape(images)[2], tf.float32)
    lam = _sample_beta_distribution(alpha, [1])[0]
    cut_w = tf.cast(w * tf.sqrt(1.0 - lam), tf.int32)
    cut_h = tf.cast(h * tf.sqrt(1.0 - lam), tf.int32)
    cx = tf.random.uniform([], minval=0, maxval=tf.cast(w, tf.int32), dtype=tf.int32)
    cy = tf.random.uniform([], minval=0, maxval=tf.cast(h, tf.int32), dtype=tf.int32)
    x1, y1 = tf.clip_by_value(cx - cut_w // 2, 0, tf.cast(w, tf.int32)), tf.clip_by_value(cy - cut_h // 2, 0, tf.cast(h, tf.int32))
    x2, y2 = tf.clip_by_value(cx + cut_w // 2, 0, tf.cast(w, tf.int32)), tf.clip_by_value(cy + cut_h // 2, 0, tf.cast(h, tf.int32))
    lam_adjusted = tf.cast(1.0 - (tf.cast((x2 - x1) * (y2 - y1), tf.float32) / (h * w + 1e-8)), base_labels.dtype)
    indices = tf.random.shuffle(tf.range(batch_size))
    shuffled_images = tf.gather(images, indices)
    y_coords, x_coords = tf.range(tf.cast(h, tf.int32))[:, None], tf.range(tf.cast(w, tf.int32))[None, :]
    in_box = (y_coords >= y1) & (y_coords < y2) & (x_coords >= x1) & (x_coords < x2)
    mask = tf.cast(~in_box, images.dtype)[None, :, :, None]
    mixed_images = images * mask + shuffled_images * (tf.cast(1.0, images.dtype) - mask)
    mixed_base = lam_adjusted * base_labels + (1.0 - lam_adjusted) * tf.gather(base_labels, indices)
    mixed_mod = lam_adjusted * mod_labels + (1.0 - lam_adjusted) * tf.gather(mod_labels, indices)
    mixed_vattu = lam_adjusted * vattu_labels + (1.0 - lam_adjusted) * tf.gather(vattu_labels, indices)
    return mixed_images, mixed_base, mixed_mod, mixed_vattu""")

# -------------------------------------------------------------
# Milestone 3: Training Pipeline and Checkpointing
# -------------------------------------------------------------
add_heading1("Milestone 3: Training Pipeline and Checkpointing")

add_paragraph(
    "Milestone 3 implements the unified two-phase training protocol, runtime extrapolation safety checks, and atomic state checkpointing to manage GPU training sessions effectively."
)

add_heading2("Activity 3.1: Two-Phase Training Protocol")

add_paragraph(
    "Step 1: Phase 1 Warmup (Epochs 1 to 5): The ImageNet pretrained backbone remains frozen. A linear learning rate schedule warms up from 0 to 1e-4, allowing the randomly initialized dense classification heads to converge without corrupting pretrained backbone weights."
)

add_paragraph(
    "Step 2: Phase 2 Fine-Tuning (Epochs 6 to 50): All backbone layers are unfrozen. Training proceeds using AdamW with weight decay 1e-4, global gradient clipping at 1.0, Exponential Moving Average (EMA momentum 0.999), and Cosine Decay learning rate scheduling down to 1e-6."
)

add_image_figure(
    "outputs/docs_diagrams/training_workflow.png",
    "Two-Phase Training Protocol and Runtime Session Budget Safety"
)

add_paragraph("Command to launch model training:")
add_terminal_command("python src/train.py --config configs/multitask_effnetv2.yaml")

add_paragraph("Configuration File: configs/multitask_effnetv2.yaml:")
add_code_block("""model:
  variant: "B0"
  weights: "imagenet"
  dropout_rate: 0.3
  img_size: 128

data:
  train_csv: "outputs/train.csv"
  val_csv: "outputs/val.csv"
  test_csv: "outputs/test.csv"
  label_maps: "outputs/label_maps.json"
  batch_size: 128
  label_smoothing: 0.1
  use_cutmix: true
  cutmix_alpha: 0.4
  cutmix_probability: 0.5
  use_augmentation: true
  rotation_degrees: 5.0
  translation_factor: 0.05
  zoom_factor: 0.05

training:
  mixed_precision: true
  learning_rate: 0.0001
  min_learning_rate: 0.000001
  weight_decay: 0.0001
  global_clipnorm: 1.0
  use_ema: true
  ema_momentum: 0.999
  warmup_epochs: 5
  finetune_epochs: 45
  early_stopping_patience: 10
  loss_weights:
    base_output: 1.0
    modifier_output: 0.5
    vattu_output: 0.5
  checkpoint_dir: "checkpoints"
  max_checkpoints_to_keep: 3""")

add_heading2("Activity 3.2: Early Timing Extrapolation and Kaggle Session Guard")

add_paragraph(
    "Step 1: Measure Epoch Duration: Cloud training environments like Kaggle and Google Colab enforce strict 12-hour session limits. TimingExtrapolationCallback tracks exact epoch durations during early training."
)

add_paragraph(
    "Step 2: Extrapolate Projected Session Duration: At the completion of Epoch 2, the callback calculates the rolling average seconds per epoch and projects the total runtime for remaining epochs. If the projected duration exceeds 11.5 hours, it raises an alert recommending batch size or epoch adjustments."
)

add_paragraph("Source Code: src/train.py (Timing Extrapolation Callback):")
add_code_block("""class TimingExtrapolationCallback(tf.keras.callbacks.Callback):
    def __init__(self, total_warmup_epochs: int, total_finetune_epochs: int, max_hours: float = 11.5, timing_epochs: int = 2):
        super().__init__()
        self.total_warmup = total_warmup_epochs
        self.total_finetune = total_finetune_epochs
        self.max_hours = max_hours
        self.timing_epochs = timing_epochs
        self.epoch_times = []

    def on_epoch_begin(self, epoch, logs=None):
        self.epoch_start = time.perf_counter()

    def on_epoch_end(self, epoch, logs=None):
        duration = time.perf_counter() - self.epoch_start
        self.epoch_times.append(duration)
        avg_sec = np.mean(self.epoch_times[-self.timing_epochs:])
        if len(self.epoch_times) == self.timing_epochs:
            total_epochs = self.total_warmup + self.total_finetune
            remaining_epochs = max(0, total_epochs - (epoch + 1))
            projected_remaining_hours = (remaining_epochs * avg_sec) / 3600.0
            logger.info("=" * 70)
            logger.info(f"KAGGLE RUNTIME BUDGET: Speed: {avg_sec:.1f}s/epoch | Remaining: {projected_remaining_hours:.2f}h | Limit: {self.max_hours:.1f}h")
            if projected_remaining_hours > self.max_hours:
                logger.warning("[ALERT] Projected training duration exceeds session budget!")
            logger.info("=" * 70)""")

add_heading2("Activity 3.3: Atomic Full-State Checkpoint Management")

add_paragraph(
    "Step 1: Multi-Asset State Serialization: Standard Keras model checkpoints save model weights only, discarding optimizer momentum, EMA variables, and iteration counts. FullStateCheckpointManager saves model weights (.weights.h5), optimizer state (.optimizer.pkl), EMA weights, and metadata JSON atomically."
)

add_paragraph(
    "Step 2: Safe Restoration: When resuming after session interruptions or during evaluation, restore_state loads the full training state, synchronizes iterations, and evaluates the best validation checkpoint."
)

add_paragraph("Source Code: src/checkpointing.py (Full State Checkpoint Manager):")
add_code_block("""class FullStateCheckpointManager:
    def __init__(self, checkpoint_dir: Union[str, Path], max_to_keep: int = 3, monitor: str = "val_loss", mode: str = "min"):
        self.checkpoint_dir = Path(checkpoint_dir)
        self.checkpoint_dir.mkdir(parents=True, exist_ok=True)
        self.max_to_keep = max_to_keep
        self.monitor = monitor
        self.mode = mode

    def save_state(self, model: tf.keras.Model, optimizer: tf.keras.optimizers.Optimizer, epoch: int, metrics: Optional[Dict[str, float]] = None, is_best: bool = False, tag: Optional[str] = None) -> Path:
        save_tag = tag if tag else (f"epoch_{epoch:03d}" if not is_best else "best_model")
        save_dir = self.checkpoint_dir / save_tag
        save_dir.mkdir(parents=True, exist_ok=True)
        model.save_weights(save_dir / "model.weights.h5")
        state_meta = {"epoch": epoch, "iterations": int(optimizer.iterations.numpy()) if hasattr(optimizer, "iterations") else 0, "monitored_metric": self.monitor, "is_best": is_best, "metrics": metrics or {}}
        with open(save_dir / "state.json", "w", encoding="utf-8") as f:
            json.dump(state_meta, f, indent=2)
        return save_dir""")

# -------------------------------------------------------------
# Milestone 4: Comprehensive Model Evaluation and Error Analysis
# -------------------------------------------------------------
add_heading1("Milestone 4: Comprehensive Model Evaluation and Error Analysis")

add_paragraph(
    "Milestone 4 presents the quantitative evaluation of the multi-head EfficientNetV2 model on the held-out test split, comparing per-head performance, recombined 630-way accuracy, baseline benchmarks, and confused character pairs."
)

add_heading2("Activity 4.1: Multi-Head and Recombined Test Set Evaluation")

add_paragraph(
    "Step 1: Per-Head Metrics: On the held-out test set, the three classification heads achieved: Base Akshara Top-1 Accuracy: 96.57%, Vowel Modifier Top-1 Accuracy: 96.60%, and Conjunct Vattu Top-1 Accuracy: 99.34%."
)

add_paragraph(
    "Step 2: Recombined 630-Way Constrained Evaluation: Passing the three predicted probability distributions through the Constrained Maximum Likelihood Decoding engine yielded a Recombined Top-1 Accuracy of 94.72% and a Recombined Top-5 Accuracy of 99.29% across all 630 dataset class directories."
)

add_image_figure(
    "outputs/docs_diagrams/evaluation_summary.png",
    "Multi-Head and Recombined Test Set Accuracy Summary"
)

add_paragraph("Command to run evaluation:")
add_terminal_command("python src/evaluate.py --test_csv outputs/test.csv --checkpoint_tag best_model")

add_paragraph("Source Code: src/evaluate.py (Evaluation Engine):")
add_code_block("""def evaluate_test_set(test_csv: str = "outputs/test.csv", label_maps_path: str = "outputs/label_maps.json", checkpoint_dir: str = "checkpoints", checkpoint_tag: str = "best_model") -> Dict[str, Any]:
    label_maps = load_label_maps(label_maps_path)
    df_test = pd.read_csv(test_csv)
    model = build_multitask_effnetv2(variant="B0", num_base=label_maps["num_base_classes"], num_mod=label_maps["num_modifier_classes"], num_vattu=label_maps["num_vattu_classes"])
    model.load_weights(f"{checkpoint_dir}/{checkpoint_tag}/model.weights.h5")
    test_ds, _, _ = create_telugu_dataset(csv_path_or_df=df_test, label_maps_or_path=label_maps, is_training=False)
    preds = model.predict(test_ds, verbose=1)
    base_probs, mod_probs, vattu_probs = parse_model_prediction_outputs(preds)
    correct_top1, correct_top5 = 0, 0
    for i in range(len(df_test)):
        rec = recombine_prediction(base_probs[i], mod_probs[i], vattu_probs[i], label_maps)
        true_cls = df_test.iloc[i]["class_name"]
        equivs = get_equivalent_classes(true_cls)
        if rec["predicted_class"] in equivs:
            correct_top1 += 1
        top5_classes = [t["class_name"] for t in rec["top_5"]]
        if any(e in top5_classes for e in equivs):
            correct_top5 += 1
    top1_acc = correct_top1 / len(df_test)
    top5_acc = correct_top5 / len(df_test)
    logger.info(f"Recombined 630-Way Test Evaluation: Top-1: {top1_acc*100:.2f}% | Top-5: {top5_acc*100:.2f}%")
    return {"top1_accuracy": top1_acc, "top5_accuracy": top5_acc}""")

add_heading2("Activity 4.2: Baseline Comparison and Confusion Analysis")

add_paragraph(
    "Step 1: Baseline Comparisons: The multi-head EfficientNetV2 model substantially outperforms traditional single-head models. A monolithic MobileNetV2 baseline trained directly on 630 flat classes achieved 85.64% Top-1 accuracy, while an Always-None majority predictor achieves 0.0% accuracy on non-trivial classes."
)

add_paragraph(
    "Step 2: Confusion Analysis: The most common remaining prediction errors occur between characters with high visual overlap, such as 'th' (థ) vs 'tha' (థ), 'dh' (ధ) vs 'dha' (ధ), and subtle diacritic loops in vowel signs."
)

# -------------------------------------------------------------
# Milestone 5: Streamlit Interactive Web Application
# -------------------------------------------------------------
add_heading1("Milestone 5: Streamlit Interactive Web Application")

add_paragraph(
    "Milestone 5 covers the development of the interactive Streamlit user interface, the canvas stroke smoothing algorithm that resolves the digital-to-physical domain shift, and real-time Telugu Unicode glyph assembly."
)

add_heading2("Activity 5.1: Two-Column Matte Technical Utility Interface")

add_paragraph(
    "Step 1: Layout and Theme Styling: The interface uses a clean two-column layout constrained to 1100px max width. A warm matte palette (#fbf9f4 background, #edebe6 surface, #a03d00 terracotta accent) and Geist typography provide a distraction-free environment."
)

add_paragraph(
    "Step 2: Responsive Columns: The Left Column houses the drawing canvas (280x280), stroke width controls (4px to 24px), clear canvas button, and recognize button. The Right Column displays the predicted Telugu Unicode character, top-3 ranked alternatives, confidence bars, and linguistic breakdown badges."
)

add_image_figure(
    "outputs/docs_diagrams/ui_layout_diagram.png",
    "Streamlit Side-by-Side Matte Technical Utility Interface Layout"
)

add_heading2("Activity 5.2: Canvas Preprocessing and Domain Shift Resolution")

add_paragraph(
    "Step 1: Canvas Gaussian Blur Smoothing: Digital HTML5 canvas drawings produce sharp aliased binary edges with a Laplacian variance of ~5545, whereas real handwritten paper scans have ink dispersion and soft edges with a Laplacian variance of ~76.7. Applying cv2.GaussianBlur(..., (0, 0), sigmaX=3.0) smooths stroke edges and closes the distribution gap."
)

add_paragraph(
    "Step 2: Telugu Unicode Synthesis: get_display_glyph synthesizes authentic multi-part Telugu graphemes by combining the base character, virama marker (\u0C4D), vattu consonant glyph, and vowel matra unicode string."
)

add_image_figure(
    "outputs/docs_diagrams/canvas_smoothing_effect.png",
    "Canvas Stroke Smoothing via Gaussian Blur for Domain Shift Resolution"
)

add_heading2("Activity 5.3: Top-3 Prediction Cards and Detailed Probability Breakdown")

add_paragraph(
    "Step 1: Primary Match Card: Displays the top predicted character in large Telugu font (44px), accompanied by the confidence percentage bar and folder class identifier."
)

add_paragraph(
    "Step 2: Component Breakdown Badges: Below the primary prediction, three badges display the individual confidence scores for Base Akshara, Vowel Modifier, and Conjunct Vattu."
)

add_paragraph("Source Code: app.py (Streamlit Application Engine):")
add_code_block("""import streamlit as st
import numpy as np
import cv2
from PIL import Image
import tensorflow as tf
from src.data.preprocessing import preprocess_image, IMAGE_SIZE
from src.data.decomposition import recombine_prediction
from src.models.multitask_effnetv2 import build_multitask_effnetv2, parse_model_prediction_outputs

MATRA_UNICODE = {"none": "", "aa": "\u0C3E", "i": "\u0C3F", "ii": "\u0C40", "u": "\u0C41", "uu": "\u0C42", "ru": "\u0C43", "ruu": "\u0C44", "e": "\u0C46", "ee": "\u0C47", "ai": "\u0C48", "o": "\u0C4A", "oo": "\u0C4B", "au": "\u0C4C", "am": "\u0C02", "ah": "\u0C03"}
VIRAMA = "\u0C4D"
VATTU_CONSONANT = {"k": "క", "kh": "ఖ", "g": "గ", "gh": "ఘ", "c": "చ", "ch": "ఛ", "j": "జ", "t": "ట", "th": "ఠ", "d": "డ", "dh": "ఢ", "n": "న", "p": "ప", "ph": "ఫ", "b": "బ", "bh": "భ", "m": "మ", "y": "య", "r": "ర", "l": "ల", "v": "వ", "s": "శ", "sh": "ష", "sa": "స", "h": "హ", "ksh": "క్ష"}

def get_display_glyph(base_char: str, mod: str, vattu: str) -> str:
    if base_char == "none":
        return f"{VIRAMA}{VATTU_CONSONANT.get(vattu, 'క')}"
    matra = MATRA_UNICODE.get(mod, "")
    if vattu != "none":
        v_char = VATTU_CONSONANT.get(vattu, "")
        return f"{base_char}{VIRAMA}{v_char}{matra}"
    return f"{base_char}{matra}"

def process_canvas_drawing(canvas_data, stroke_width):
    alpha = canvas_data[:, :, 3]
    if not np.any(alpha > 20):
        return None
    gray = 255 - alpha
    coords = cv2.findNonZero((gray < 240).astype(np.uint8))
    x, y, w, h = cv2.boundingRect(coords)
    pad = int(max(w, h) * 0.15)
    crop = gray[max(0, y - pad):min(gray.shape[0], y + h + pad), max(0, x - pad):min(gray.shape[1], x + w + pad)]
    smoothed = cv2.GaussianBlur(crop, (0, 0), sigmaX=3.0)
    return preprocess_image(smoothed, img_size=IMAGE_SIZE)""")

# -------------------------------------------------------------
# Milestone 6: Deployment and Verification
# -------------------------------------------------------------
add_heading1("Milestone 6: Deployment and Verification")

add_paragraph(
    "Milestone 6 provides complete instructions for local application deployment, verification across drawing and image upload modes, and public deployment using Ngrok tunneling."
)

add_heading2("Activity 6.1: Preparing the Application for Local Deployment")

add_paragraph("Step 1: Set Up Python Virtual Environment:")
add_terminal_command("python -m venv venv\nsource venv/bin/activate  # On macOS/Linux\n# or: .\\venv\\Scripts\\activate  # On Windows")

add_paragraph("Step 2: Install Required Dependencies:")
add_terminal_command("pip install -r requirements.txt")

add_paragraph("Complete requirements.txt listing:")
add_code_block("""tensorflow>=2.16.0
keras>=3.0.0
numpy>=1.24.0
pandas>=2.0.0
scikit-learn>=1.3.0
pillow>=10.0.0
pyyaml>=6.0.0
tqdm>=4.65.0
streamlit>=1.30.0
streamlit-drawable-canvas>=0.9.0""")

add_heading2("Activity 6.2: Local Testing and Verification")

add_paragraph("Step 1: Start the Streamlit Application Server:")
add_terminal_command("streamlit run app.py --server.port 8501")

add_paragraph(
    "Step 2: Verify Recognition Pipeline: Open http://localhost:8501 in a web browser. Test character recognition using both the interactive canvas drawing tool and uploaded test image files."
)

add_heading2("Activity 6.3: Public Deployment via Ngrok")

add_paragraph(
    "Step 1: Install and Configure Ngrok: Ngrok creates a secure HTTPS tunnel to the local Streamlit port, allowing public access without server configuration."
)
add_terminal_command("pip install pyngrok\nngrok config add-authtoken YOUR_NGROK_AUTHTOKEN")

add_paragraph("Step 2: Run the Public Deployment Script (run_public.py):")
add_terminal_command("python run_public.py")

add_paragraph("Source Code: run_public.py (Public Tunnel Deployment):")
add_code_block("""import os
import sys
import subprocess
from pyngrok import ngrok

def start_public_app():
    port = 8501
    public_url = ngrok.connect(port, proto="http")
    print("=" * 60)
    print("TELUGU HCR V4 WEB APP IS LIVE!")
    print(f"Public URL: {public_url}")
    print("=" * 60)
    cmd = [sys.executable, "-m", "streamlit", "run", "app.py", f"--server.port={port}", "--server.headless=true"]
    subprocess.run(cmd)

if __name__ == "__main__":
    start_public_app()""")

add_heading2("Important Notes:")
add_bullet("Apple Silicon Acceleration: On macOS with Apple Silicon, TensorFlow utilizes Metal Performance Shaders (MPS). In automated test environments, CPU execution is enforced to prevent GPU kernel locks during rapid evaluation cycles.", bold_prefix="Hardware Acceleration: ")
add_bullet("Canvas Stroke Thickness: For optimal recognition accuracy, stroke width should be set between 8px and 16px, matching the pen line thickness of standard dataset images.", bold_prefix="Drawing Technique: ")
add_bullet("Ngrok Tunnel Expiration: On free Ngrok accounts, public tunnel URLs expire after session termination. Restarting run_public.py assigns a new active URL.", bold_prefix="Tunnel Persistence: ")

# -------------------------------------------------------------
# Exploring the Web Application
# -------------------------------------------------------------
add_heading1("Exploring the Web Application:")

add_heading2("Home / Recognition Page")
add_paragraph(
    "The main application page features a top navigation bar with the brand title 'Telugu Akshara Recognizer', model status indicator (Online 94.7%), and active parameter badge (EfficientNetV2-B0 5.9M). The page is organized in a balanced two-column utility layout."
)

add_heading2("Input Canvas and Drawing Controls")
add_paragraph(
    "The left column provides an HTML5 drawing canvas with real-time mouse and stylus tracking. Users can select stroke widths from 4px to 24px, switch between drawing mode and image upload mode, and trigger recognition via the 'Recognize Character' action button."
)

add_heading2("Recognition Results and Component Breakdown")
add_paragraph(
    "The right column presents the recognition results. The primary prediction card highlights the recognized character in large Telugu typography, accompanied by the recombined confidence percentage. Three component badges display the breakdown for Base Akshara, Vowel Modifier, and Conjunct Vattu. Below the primary card, alternative Top-2 and Top-3 candidate cards provide secondary predictions with confidence bars."
)

# -------------------------------------------------------------
# Conclusion
# -------------------------------------------------------------
add_heading1("Conclusion")

add_paragraph(
    "Telugu HCR v4 resolves the challenge of handwritten Telugu compound character recognition through linguistic decomposition and multi-task deep learning. By decomposing 630 character classes into 52 base aksharas, 16 vowel modifiers, and 37 conjunct vattus, the EfficientNetV2-B0 model achieves 94.72% Top-1 and 99.29% Top-5 recombined accuracy. The single source of truth preprocessing pipeline and Gaussian blur stroke smoothing effectively bridge the domain gap between digital canvas drawings and physical paper scans. Future work includes extending the pipeline to writer-independent grouping splits, sentence-level connected text recognition, and edge deployment via TensorFlow Lite."
)

# Save DOCX
doc.save(str(DOCX_OUT))
print(f"Saved DOCX to {DOCX_OUT}")

# Copy to Downloads
shutil.copy(str(DOCX_OUT), str(DOWNLOADS_DIR / "Telugu_HCR_v4_Project_Documentation.docx"))
print(f"Copied DOCX to {DOWNLOADS_DIR / 'Telugu_HCR_v4_Project_Documentation.docx'}")

# Generate Markdown version
md_lines = []
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    if p.runs and p.runs[0].font.size and p.runs[0].font.size.pt == 20:
        md_lines.append(f"# {txt}\n")
    elif p.runs and p.runs[0].font.size and p.runs[0].font.size.pt == 15:
        md_lines.append(f"\n## {txt}\n")
    elif p.runs and p.runs[0].font.size and p.runs[0].font.size.pt == 13:
        md_lines.append(f"\n### {txt}\n")
    elif p.runs and p.runs[0].font.size and p.runs[0].font.size.pt == 11.5:
        md_lines.append(f"\n#### {txt}\n")
    elif p.style.name == 'List Bullet':
        md_lines.append(f"- {txt}")
    else:
        md_lines.append(f"{txt}\n")

with open(MD_OUT, "w", encoding="utf-8") as f:
    f.write("\n".join(md_lines))
print(f"Saved Markdown to {MD_OUT}")
